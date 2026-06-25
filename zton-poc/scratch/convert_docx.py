import docx
import os

def docx_to_markdown(docx_path, md_path):
    doc = docx.Document(docx_path)
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
            
        # Basic mapping of styles to Markdown headings
        if para.style.name.startswith("Heading 1"):
            md_lines.append(f"# {text}\n")
        elif para.style.name.startswith("Heading 2"):
            md_lines.append(f"## {text}\n")
        elif para.style.name.startswith("Heading 3"):
            md_lines.append(f"### {text}\n")
        elif para.style.name.startswith("Heading 4"):
            md_lines.append(f"#### {text}\n")
        elif para.style.name.startswith("List Bullet"):
            md_lines.append(f"* {text}")
        elif para.style.name.startswith("List Number"):
            md_lines.append(f"1. {text}")
        else:
            md_lines.append(text + "\n")
            
    # Process tables if there are any
    if doc.tables:
        md_lines.append("\n--- TABLES FOUND IN DOCUMENT ---\n")
        for i, table in enumerate(doc.tables):
            md_lines.append(f"\n### Table {i+1}\n")
            for row in table.rows:
                row_cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(row_cells) + " |")
                
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

if __name__ == "__main__":
    # Let's extract the docx file
    docx_file = r"c:\Users\rishi\Documents\NPS-LAB-EL\NPS-LAB-EL\NPS report (1).docx"
    md_file = r"c:\Users\rishi\Documents\NPS-LAB-EL\NPS-LAB-EL\full_complete_report.md"
    
    if os.path.exists(docx_file):
        print(f"Converting {docx_file} to {md_file}...")
        docx_to_markdown(docx_file, md_file)
        print("Conversion complete!")
    else:
        print(f"Error: {docx_file} not found.")
